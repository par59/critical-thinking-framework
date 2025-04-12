# For image to word 

from PIL import Image
import pytesseract
from docx import Document

# Set Tesseract path (adjust if installed somewhere else)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Load your image
image_path = r"D:\indiannn\lovely\image\700.jpg"
image = Image.open(image_path)

# Extract text from the image
extracted_text = pytesseract.image_to_string(image)

# Create a Word document
doc = Document()
doc.add_heading("Extracted Text from Image", level=1)
doc.add_paragraph(extracted_text)

# Save the document
doc_path = r"D:\indiannn\lovely\image\only.docx"
doc.save(doc_path)

print(f"✅ Text successfully extracted and saved to: {doc_path}")