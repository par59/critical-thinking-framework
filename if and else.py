from PIL import Image
import pytesseract

# Set the path to Tesseract executable (adjust if yours is different)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Load your image (use full path or place image in same folder)
image_path = r"C:\Users\paria\OneDrive\Desktop\indiannn\lovely\image\66.jpg"
image = Image.open(image_path)

# Extract text from image
extracted_text = pytesseract.image_to_string(image)

# Print the extracted text
print("Extracted Text:\n")
print(extracted_text)

# Optional: Save to a text file
with open("output_text.txt", "w", encoding="utf-8") as file:
    file.write(extracted_text)
from PIL import Image
import pytesseract
from docx import Document

# Set Tesseract path (adjust if installed somewhere else)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Load your image
image_path = r"C:\Users\paria\OneDrive\Desktop\indiannn\lovely\image\66.jpg"
image = Image.open(image_path)

# Extract text from the image
extracted_text = pytesseract.image_to_string(image)

# Create a Word document
doc = Document()
doc.add_heading("Extracted Text from Image", level=1)
doc.add_paragraph(extracted_text)

# Save the document
doc_path = r"C:\Users\paria\OneDrive\Desktop\indiannn\lovely\image\ExtractedText.docx"
doc.save(doc_path)

print(f"✅ Text successfully extracted and saved to: {doc_path}")
