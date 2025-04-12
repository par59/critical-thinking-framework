from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
from docx import Document

# YouTube video ID
video_id = "AKiGG8k-FEM"

# Create Word document
doc = Document()
doc.add_heading('Translated YouTube Transcript (Hindi ➜ English)', 0)

try:
    # Get available transcripts
    transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

    # Find Hindi transcript and translate to English
    transcript = transcript_list.find_generated_transcript(['hi']).translate('en').fetch()

    # Format transcript as plain text
    formatter = TextFormatter()
    formatted_text = formatter.format_transcript(transcript)

    # Save to .docx
    doc.add_paragraph(formatted_text)
    doc.save("YouTube_Transcript_Translated.docx")

    print("✅ Transcript translated from Hindi to English and saved as 'YouTube_Transcript_Translated111.docx'.")

except Exception as e:
    print("❌ Error:", e)
