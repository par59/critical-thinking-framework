import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from docx import Document

# Path to Tesseract OCR executable
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ✅ Updated file paths
pdf_path = r"D:\indiannn\lovely\30th Aug 2024 Shift 1.pdf"
doc_path = r"D:\indiannn\lovely\shift 1.docx"

# Create a new Word document
doc = Document()
doc.add_heading("Extracted Text from PDF", level=1)

# Open the PDF
pdf_file = fitz.open(pdf_path)

# Loop through all pages
for page_number in range(len(pdf_file)):
    page = pdf_file[page_number]
    
    # Convert page to image
    pix = page.get_pixmap(dpi=300)
    img_data = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_data))
    
    # OCR: Extract text from image
    text = pytesseract.image_to_string(img)
    
    # Add text to Word file
    doc.add_heading(f"Page {page_number + 1}", level=2)
    doc.add_paragraph(text)

# Save the Word document
doc.save(doc_path)
print(f"✅ Success! Text saved to: {doc_path}")
